"""
Модуль для экспорта статистики в CSV и Excel форматы
"""
import csv
import os
import re
import time
from collections import Counter
from datetime import datetime
from typing import List, Dict
from pathlib import Path

from requests import post
from core.logging_config import logger

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, Reference
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from core.employee_tagger import normalize_person_name

POST_EXPORT_COLUMNS = [
    'ID', 'Дата', 'Текст', 'Ссылка', 'Автор',
    'Лайки', 'Комментарии', 'Репосты', 'Популярность',
]

# Диаграммы справа от таблицы (колонка K), фиксированные строки — не зависят от числа постов
POST_CHART_ANCHORS = ('K2', 'K22', 'K42', 'K62')
POST_CHART_SIZE = {'height': 7.5, 'width': 14}

TEACHER_STATS_WORD_COLUMNS = ['Преподаватель', 'Количество постов']
TEACHER_POSTS_WORD_COLUMNS = [
    'Текст поста', 'Ссылка', 'ФИО преподавателя', 'Хэштег преподавателя', 'Хэштег кафедры',
]
TEACHER_POST_TEXT_WORD_LIMIT = 15


_MEDIA_TYPE_RU = {
    'photo': 'Фото',
    'video': 'Видео',
    'clip': 'Клип',
}


def _media_type_ru(value: str) -> str:
    if not value:
        return ''
    return _MEDIA_TYPE_RU.get(str(value).lower(), value)


class StatisticsExporter:
    """Класс для экспорта статистики в различные форматы"""
    def __init__(self, export_dir: str = './exports'):
        self.export_dir = export_dir

    @staticmethod
    def _post_export_row(post: Dict) -> Dict:
        likes = post.get('likes', 0) or 0
        comments = post.get('comments', 0) or 0
        shares = post.get('shares', 0) or 0
        popularity = post.get('popularity', likes + comments + shares)
        return {
            'ID': post.get('post_id', ''),
            'Дата': post.get('date', ''),
            'Текст': post.get('text', ''),
            'Ссылка': post.get('post_url', ''),
            'Автор': post.get('author_name', ''),
            'Лайки': likes,
            'Комментарии': comments,
            'Репосты': shares,
            'Популярность': popularity,
        }

    def _ensure_export_dir(self):
        if not os.path.exists(self.export_dir):
            Path(self.export_dir).mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _add_posts_summary_rows(ws, last_data_row: int) -> tuple[int, int]:
        """Строки ИТОГО и СРЕДНЕЕ сразу под данными."""
        total_row = last_data_row + 2
        avg_row = last_data_row + 3
        total_fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
        avg_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')

        ws.merge_cells(f'A{total_row}:E{total_row}')
        ws[f'A{total_row}'] = 'ИТОГО'
        ws[f'A{total_row}'].font = Font(bold=True, size=11)
        ws[f'A{total_row}'].alignment = Alignment(horizontal='left', vertical='center')
        ws[f'A{total_row}'].fill = total_fill

        for col in ('F', 'G', 'H', 'I'):
            ws[f'{col}{total_row}'] = f'=SUM({col}2:{col}{last_data_row})'
            ws[f'{col}{total_row}'].fill = total_fill
            ws[f'{col}{total_row}'].font = Font(bold=True)
            ws[f'{col}{total_row}'].alignment = Alignment(horizontal='center')

        ws.merge_cells(f'A{avg_row}:E{avg_row}')
        ws[f'A{avg_row}'] = 'СРЕДНЕЕ'
        ws[f'A{avg_row}'].font = Font(bold=True, size=10, color='666666')
        ws[f'A{avg_row}'].alignment = Alignment(horizontal='left', vertical='center')
        ws[f'A{avg_row}'].fill = avg_fill

        for col in ('F', 'G', 'H', 'I'):
            ws[f'{col}{avg_row}'] = f'=AVERAGE({col}2:{col}{last_data_row})'
            ws[f'{col}{avg_row}'].fill = avg_fill
            ws[f'{col}{avg_row}'].font = Font(size=10)
            ws[f'{col}{avg_row}'].alignment = Alignment(horizontal='center')
            ws[f'{col}{avg_row}'].number_format = '0.00'

        return total_row, avg_row

    @staticmethod
    def _add_posts_charts(ws, posts_count: int) -> None:
        """Диаграммы на том же листе, справа от таблицы."""
        if posts_count <= 0:
            return

        top_n = min(10, posts_count)
        trend_n = min(20, posts_count)
        h, w = POST_CHART_SIZE['height'], POST_CHART_SIZE['width']

        chart1 = BarChart()
        chart1.type = 'col'
        chart1.grouping = 'clustered'
        chart1.title = f'Топ {top_n} постов по популярности'
        chart1.y_axis.title = 'Популярность'
        chart1.x_axis.title = 'Дата'
        chart1.height = h
        chart1.width = w
        chart1.add_data(
            Reference(ws, min_col=9, min_row=1, max_row=top_n + 1),
            titles_from_data=True,
        )
        chart1.set_categories(Reference(ws, min_col=2, min_row=2, max_row=top_n + 1))
        ws.add_chart(chart1, POST_CHART_ANCHORS[0])

        cats = Reference(ws, min_col=2, min_row=2, max_row=trend_n + 1)

        chart2 = LineChart()
        chart2.title = 'Тренд лайков (первые 20 постов)'
        chart2.y_axis.title = 'Лайки'
        chart2.x_axis.title = 'Дата'
        chart2.height = h
        chart2.width = w
        chart2.add_data(
            Reference(ws, min_col=6, min_row=1, max_row=trend_n + 1),
            titles_from_data=True,
        )
        chart2.set_categories(cats)
        ws.add_chart(chart2, POST_CHART_ANCHORS[1])

        chart3 = LineChart()
        chart3.title = 'Тренд комментариев (первые 20 постов)'
        chart3.y_axis.title = 'Комментарии'
        chart3.height = h
        chart3.width = w
        chart3.add_data(
            Reference(ws, min_col=7, min_row=1, max_row=trend_n + 1),
            titles_from_data=True,
        )
        chart3.set_categories(cats)
        ws.add_chart(chart3, POST_CHART_ANCHORS[2])

        chart4 = LineChart()
        chart4.title = 'Тренд репостов (первые 20 постов)'
        chart4.y_axis.title = 'Репосты'
        chart4.height = h
        chart4.width = w
        chart4.add_data(
            Reference(ws, min_col=8, min_row=1, max_row=trend_n + 1),
            titles_from_data=True,
        )
        chart4.set_categories(cats)
        ws.add_chart(chart4, POST_CHART_ANCHORS[3])

    # ===== ЭКСПОРТ В CSV =====
    def export_posts_to_csv(self, posts: List[Dict], filename: str = None) -> str:
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'экспорт_постов_{timestamp}.csv'
        
        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            with open(filepath, 'w', newline='', encoding='utf-8-sig') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=POST_EXPORT_COLUMNS)
                writer.writeheader()
                for post in posts:
                    writer.writerow(self._post_export_row(post))
            logger.info("Посты экспортированы в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта CSV: %s", e, exc_info=True)
            return None

    @staticmethod
    def _truncate_words(text: str, limit: int = TEACHER_POST_TEXT_WORD_LIMIT) -> str:
        clean = re.sub(r'#[\wА-Яа-яЁё]+', '', text or '')
        clean = re.sub(r'\s+', ' ', clean).strip()
        words = clean.split()
        if len(words) <= limit:
            return ' '.join(words)
        return ' '.join(words[:limit]) + '…'

    @staticmethod
    def _format_teacher_name(row: Dict) -> str:
        author = row.get('author_name', '') or ''
        if author and not str(author).startswith('#'):
            return normalize_person_name(str(author))
        return author

    @staticmethod
    def _format_teacher_post_row(row: Dict) -> List[str]:
        return [
            StatisticsExporter._truncate_words(str(row.get('text', '') or '')),
            str(row.get('post_url', '') or ''),
            StatisticsExporter._format_teacher_name(row),
            str(row.get('teacher_hashtag', '') or ''),
            str(row.get('department_hashtag', '') or ''),
        ]

    @staticmethod
    def _build_teacher_post_stats(rows: List[Dict]) -> List[tuple[str, int]]:
        counter = Counter()
        for row in rows:
            name = StatisticsExporter._format_teacher_name(row) or 'Неизвестно'
            counter[name] += 1
        return sorted(counter.items(), key=lambda item: (-item[1], item[0]))

    @staticmethod
    def _shade_word_cell(cell, fill_hex: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), fill_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _style_word_table_header(table, headers: List[str], fill_hex: str = '4472C4') -> None:
        header_cells = table.rows[0].cells
        for idx, title_text in enumerate(headers):
            cell = header_cells[idx]
            cell.text = ''
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(title_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            StatisticsExporter._shade_word_cell(cell, fill_hex)

    @staticmethod
    def _fill_word_table_row(cells, values: List[str], *, center_cols: set = None, font_size: int = 9, fill_hex: str = None) -> None:
        center_cols = center_cols or set()
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.text = ''
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(str(value))
            run.font.size = Pt(font_size)
            if fill_hex:
                StatisticsExporter._shade_word_cell(cell, fill_hex)

    def export_teachers_posts_to_word(self, rows: List[Dict], period_label: str = '', filename: str = None) -> str:
        if not DOCX_AVAILABLE:
            logger.info("python-docx не установлен. Экспорт в Word недоступен.")
            return None

        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'посты_преподавателей_{timestamp}.docx'

        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            doc = Document()
            try:
                doc.styles['Normal'].font.name = 'Times New Roman'
                doc.core_properties.title = 'Посты преподавателей'
                doc.core_properties.author = 'Пресс-центр ЧИБГУ'
            except Exception:
                pass
            section = doc.sections[0]
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)

            title = doc.add_heading('Отчёт по преподавателям', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if period_label:
                period_para = doc.add_paragraph(period_label)
                period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                period_para.runs[0].font.size = Pt(11)
                period_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            if not rows:
                doc.add_paragraph('За выбранный период постов не найдено.')
                doc.save(filepath)
                return filepath

            posts_heading = doc.add_heading('Посты преподавателей', level=2)
            posts_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            posts_table = doc.add_table(rows=1, cols=len(TEACHER_POSTS_WORD_COLUMNS))
            posts_table.style = 'Table Grid'
            posts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._style_word_table_header(posts_table, TEACHER_POSTS_WORD_COLUMNS, fill_hex='4472C4')

            for row_idx, row_data in enumerate(rows, start=1):
                values = self._format_teacher_post_row(row_data)
                row_fill = 'F2F2F2' if row_idx % 2 == 0 else None
                self._fill_word_table_row(
                    posts_table.add_row().cells,
                    values,
                    fill_hex=row_fill,
                )

            for idx, width in enumerate((Cm(7), Cm(4.5), Cm(4), Cm(3.5), Cm(4.5))):
                posts_table.columns[idx].width = width

            doc.add_paragraph()
            stats_heading = doc.add_heading('Статистика по преподавателям', level=2)
            stats_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            stats_table = doc.add_table(rows=1, cols=len(TEACHER_STATS_WORD_COLUMNS))
            stats_table.style = 'Table Grid'
            stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._style_word_table_header(stats_table, TEACHER_STATS_WORD_COLUMNS, fill_hex='2E75B6')

            teacher_stats = self._build_teacher_post_stats(rows)
            for idx, (name, count) in enumerate(teacher_stats, start=1):
                row_fill = 'F2F2F2' if idx % 2 == 0 else None
                self._fill_word_table_row(
                    stats_table.add_row().cells,
                    [name, str(count)],
                    center_cols={1},
                    fill_hex=row_fill,
                )

            for idx, width in enumerate((Cm(9), Cm(4))):
                stats_table.columns[idx].width = width

            doc.add_paragraph()
            footer = doc.add_paragraph(f'Всего постов: {len(rows)}')
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            doc.save(filepath)
            logger.info("Посты преподавателей экспортированы в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта Word: %s", e, exc_info=True)
            return None

    def export_employees_to_csv(self, employees: List[Dict], filename: str = None) -> str:
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'экспорт_преподавателей_{timestamp}.csv'
        
        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            with open(filepath, 'w', newline='', encoding='utf-8-sig') as csvfile:
                fieldnames = ['ФИО', 'Упоминаний', 'Постов', 'Всего лайков',
                              'Средние лайки на пост']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                for emp in employees:
                    writer.writerow({
                        'ФИО': emp.get('employee', ''),
                        'Упоминаний': emp.get('mention_count', 0),
                        'Постов': emp.get('post_count', 0),
                        'Всего лайков': emp.get('total_likes', 0),
                        'Средние лайки на пост': f"{emp.get('avg_post_likes', 0):.2f}"
                    })
            logger.info("Преподаватели экспортированы в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта CSV: %s", e, exc_info=True)
            return None

    def export_media_to_csv(self, media: List[Dict], filename: str = None) -> str:
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'экспорт_медиа_{timestamp}.csv'
        
        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            with open(filepath, 'w', newline='', encoding='utf-8-sig') as csvfile:
                fieldnames = ['Ключ медиа', 'Тип', 'Лайки', 'Комментарии', 'Дата']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                for m in media:
                    writer.writerow({
                        'Ключ медиа': m.get('media_key', ''),
                        'Тип': _media_type_ru(m.get('type', '')),
                        'Лайки': m.get('likes', m.get('metric_value', 0)),
                        'Комментарии': m.get('comments', 0),
                        'Дата': m.get('date', '')
                    })
            logger.info("Медиа экспортировано в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта CSV: %s", e, exc_info=True)
            return None

    # ===== ЭКСПОРТ В EXCEL =====
    def export_posts_to_excel(self, posts: List[Dict], filename: str = None) -> str:
        if not OPENPYXL_AVAILABLE:
            logger.info("openpyxl не установлен. Использую CSV-фоллбек.")
            return self.export_posts_to_csv(posts, filename.replace('.xlsx', '.csv') if filename else None)

        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'экспорт_постов_{timestamp}.xlsx'

        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            wb = openpyxl.Workbook()
            wb.properties.creator = 'Пресс-центр ЧИБГУ'
            wb.properties.title = 'Экспорт постов'
            ws = wb.active
            ws.title = 'Посты'

            ws.append(POST_EXPORT_COLUMNS)

            header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFF', size=11)
            thin = Side(style='thin', color='D9D9D9')
            header_border = Border(left=thin, right=thin, top=thin, bottom=thin)
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = header_border

            zero_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
            positive_fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
            zero_font = Font(color='9C0006')
            positive_font = Font(color='006100')

            last_data_row = 1
            for post in posts:
                row_data = self._post_export_row(post)
                ws.append([row_data[col] for col in POST_EXPORT_COLUMNS])
                last_data_row += 1
                row_idx = last_data_row

                ws[f'E{row_idx}'].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

                for col_letter, key in [('F', 'Лайки'), ('G', 'Комментарии'), ('H', 'Репосты'), ('I', 'Популярность')]:
                    value = row_data[key]
                    cell = ws[f'{col_letter}{row_idx}']
                    if value == 0:
                        cell.fill = zero_fill
                        cell.font = zero_font
                    else:
                        cell.fill = positive_fill
                        cell.font = positive_font

            if len(posts) > 0:
                self._add_posts_summary_rows(ws, last_data_row)

            ws.column_dimensions['A'].width = 10
            ws.column_dimensions['B'].width = 18
            ws.column_dimensions['C'].width = 48
            ws.column_dimensions['D'].width = 36
            ws.column_dimensions['E'].width = 34
            ws.column_dimensions['F'].width = 11
            ws.column_dimensions['G'].width = 14
            ws.column_dimensions['H'].width = 11
            ws.column_dimensions['I'].width = 14
            ws.column_dimensions['K'].width = 2

            for row in ws.iter_rows(min_row=2, min_col=3, max_col=5, max_row=last_data_row):
                for cell in row:
                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

            for row in ws.iter_rows(min_row=2, min_col=6, max_col=9, max_row=last_data_row):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center', vertical='center')

            ws.freeze_panes = 'A2'
            if len(posts) > 0:
                ws.auto_filter.ref = f'A1:I{last_data_row}'
                self._add_posts_charts(ws, len(posts))

            wb.save(filepath)
            logger.info("Посты экспортированы в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта Excel: %s", e, exc_info=True)
            return None

    def export_employees_to_excel(self, employees: List[Dict], filename: str = None) -> str:
        if not OPENPYXL_AVAILABLE:
            return self.export_employees_to_csv(employees, filename.replace('.xlsx', '.csv') if filename else None)
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'экспорт_преподавателей_{timestamp}.xlsx'
        
        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'Преподаватели'
            
            headers = ['ФИО', 'Упоминаний', 'Постов', 'Всего лайков', 'Средние лайки на пост']
            ws.append(headers)
            
            # Форматирование заголовков
            header_fill = PatternFill(start_color='70AD47', end_color='70AD47', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFF', size=11)
            header_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                  top=Side(style='thin'), bottom=Side(style='thin'))
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = header_border
            
            # Добавление данных преподавателей
            for emp in employees:
                mention_count = emp.get('mention_count', 0) or 0
                post_count = emp.get('post_count', 0) or 0
                total_likes = emp.get('total_likes', 0) or 0
                avg_likes = emp.get('avg_post_likes', 0) or 0
                
                ws.append([
                    emp.get('employee', ''),
                    mention_count,
                    post_count,
                    total_likes,
                    f"{avg_likes:.2f}"
                ])
            
            # Итоговая строка
            if len(employees) > 0:
                total_row = len(employees) + 2
                ws[f'A{total_row}'] = 'ИТОГО'
                ws[f'A{total_row}'].font = Font(bold=True, size=11)
                ws[f'A{total_row}'].fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
                
                # Формулы
                ws[f'B{total_row}'] = f'=SUM(B2:B{total_row-1})'
                ws[f'C{total_row}'] = f'=SUM(C2:C{total_row-1})'
                ws[f'D{total_row}'] = f'=SUM(D2:D{total_row-1})'
                ws[f'E{total_row}'] = f'=AVERAGE(E2:E{total_row-1})'
                
                for col in ['B', 'C', 'D', 'E']:
                    ws[f'{col}{total_row}'].fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
                    ws[f'{col}{total_row}'].font = Font(bold=True)
            
            # Ширина колонок
            ws.column_dimensions['A'].width = 30
            ws.column_dimensions['B'].width = 14
            ws.column_dimensions['C'].width = 12
            ws.column_dimensions['D'].width = 15
            ws.column_dimensions['E'].width = 18
            
            # Центрирование
            for row in ws.iter_rows(min_row=2, min_col=2, max_col=5, max_row=len(employees)+1):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center')
            
            # Freeze pane
            ws.freeze_panes = 'A2'
            
            # Автофильтр
            if len(employees) > 0:
                ws.auto_filter.ref = f'A1:E{len(employees)+1}'
                
                # Диаграмма: топ преподавателей по упоминаниям
                top_n = min(10, len(employees))
                chart = BarChart()
                chart.type = 'col'
                chart.title = f'Топ {top_n} преподавателей по упоминаниям'
                chart.y_axis.title = 'Количество упоминаний'
                chart.x_axis.title = 'ФИО'
                
                data = Reference(ws, min_col=2, min_row=1, max_row=top_n+1)
                cats = Reference(ws, min_col=1, min_row=2, max_row=top_n+1)
                chart.add_data(data, titles_from_data=True)
                chart.set_categories(cats)
                chart.height = 10
                chart.width = 20
                ws.add_chart(chart, 'A' + str(len(employees) + 4))
            
            wb.save(filepath)
            logger.info("Преподаватели экспортированы в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта Excel: %s", e, exc_info=True)
            return None

    def export_employee_dictionary_to_excel(self, employees: List[Dict], filename: str = None) -> str:
        if not OPENPYXL_AVAILABLE:
            logger.info("openpyxl не установлен. Использую CSV-фоллбек для словаря преподавателей.")
            return self.export_employees_to_csv(employees, filename.replace('.xlsx', '.csv') if filename else None)

        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'справочник_преподавателей_{timestamp}.xlsx'

        self._ensure_export_dir()
        self._cleanup_old_dictionary_files(prefix='справочник_преподавателей_')

        filepath = os.path.join(self.export_dir, filename)
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'Справочник преподавателей'

            headers = ['Кафедра', 'Хэштег кафедры', 'ФИО преподавателя', 'Хэштег преподавателя', 'Источник URL']
            ws.append(headers)

            header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFF', size=11)
            header_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                  top=Side(style='thin'), bottom=Side(style='thin'))
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = header_border

            for emp in employees:
                ws.append([
                    emp.get('department_name', ''),
                    emp.get('department_hashtag', ''),
                    emp.get('full_name', ''),
                    emp.get('hashtag', ''),
                    emp.get('source_url', '')
                ])

            ws.column_dimensions['A'].width = 30
            ws.column_dimensions['B'].width = 24
            ws.column_dimensions['C'].width = 26
            ws.column_dimensions['D'].width = 24
            ws.column_dimensions['E'].width = 40

            for row in ws.iter_rows(min_row=2, min_col=1, max_col=5, max_row=len(employees) + 1):
                for cell in row:
                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

            ws.freeze_panes = 'A2'
            if len(employees) > 0:
                ws.auto_filter.ref = f'A1:E{len(employees)+1}'

            wb.save(filepath)
            logger.info("Справочник преподавателей экспортирован в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта словаря Excel: %s", e, exc_info=True)
            return None

    def _cleanup_old_dictionary_files(self, prefix='employee_dictionary_'):
        try:
            for filename in os.listdir(self.export_dir):
                if filename.startswith(prefix) and filename.endswith('.xlsx'):
                    os.remove(os.path.join(self.export_dir, filename))
        except Exception:
            pass

    def export_comprehensive_report(self, stats: Dict, filename: str = None) -> str:
        if not OPENPYXL_AVAILABLE:
            logger.info("openpyxl не установлен. Экспорт отчета недоступен.")
            return None
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'сводный_отчёт_{timestamp}.xlsx'
        
        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'Сводка'
            
            ws['A1'] = 'СТАТИСТИКА АРХИВА'
            ws['A1'].font = Font(size=16, bold=True)
            ws.merge_cells('A1:B1')
            
            ws['A3'] = 'Всего постов:'; ws['B3'] = stats.get('total_posts', 0)
            ws['A4'] = 'Всего фото:'; ws['B4'] = stats.get('total_photos', 0)
            ws['A5'] = 'Всего видео:'; ws['B5'] = stats.get('total_videos', 0)
            ws['A6'] = 'Всего лайков:'; ws['B6'] = stats.get('total_likes', 0)
            ws['A7'] = 'Средние лайки на пост:'; ws['B7'] = f"{stats.get('avg_likes_per_post', 0):.2f}"
            
            ws.column_dimensions['A'].width = 25
            ws.column_dimensions['B'].width = 20
            
            logger.info("Комплексный отчет создан: %s", filepath)
            wb.save(filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта отчета: %s", e, exc_info=True)
            return None

    # ===== УТИЛИТЫ =====
    def get_export_files(self) -> List[str]:
        if not os.path.exists(self.export_dir):
            return []
        return os.listdir(self.export_dir)

    def clean_old_exports(self, max_age_days: int = 30) -> int:
        removed_count = 0
        max_age_seconds = max_age_days * 86400
        try:
            for filename in os.listdir(self.export_dir):
                filepath = os.path.join(self.export_dir, filename)
                file_age = time.time() - os.path.getmtime(filepath)
                if file_age > max_age_seconds:
                    os.remove(filepath)
                    removed_count += 1
            logger.info("Удалено старых экспортов: %d", removed_count)
        except Exception as e:
            logger.error("Ошибка при удалении старых файлов: %s", e, exc_info=True)
        return removed_count
    
    def export_media_to_excel(self, media: List[Dict], filename: str = None) -> str:
        if not OPENPYXL_AVAILABLE:
            return self.export_media_to_csv(media, filename.replace('.xlsx', '.csv') if filename else None)
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'экспорт_медиа_{timestamp}.xlsx'
        
        self._ensure_export_dir()
        filepath = os.path.join(self.export_dir, filename)
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'Медиа'
            
            headers = ['ID поста', 'Ключ медиа', 'Тип', 'Лайки', 'Комментарии', 'Репосты', 'Дата']
            ws.append(headers)
            
            # Форматирование заголовков
            header_fill = PatternFill(start_color='9C27B0', end_color='9C27B0', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFF', size=11)
            header_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                  top=Side(style='thin'), bottom=Side(style='thin'))
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = header_border
            
            # Добавление данных медиа
            for m in media:
                likes = m.get('likes', 0) or 0
                comments = m.get('comments', 0) or 0
                shares = m.get('shares', 0) or 0
                
                ws.append([
                    m.get('post_id', ''),
                    m.get('media_key', ''),
                    _media_type_ru(m.get('media_type', '')),
                    likes,
                    comments,
                    shares,
                    m.get('date', '')
                ])
            
            # Итоговая строка
            if len(media) > 0:
                total_row = len(media) + 2
                ws[f'A{total_row}'] = 'ИТОГО'
                ws[f'A{total_row}'].font = Font(bold=True, size=11)
                ws[f'A{total_row}'].fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
                
                # Формулы
                ws[f'D{total_row}'] = f'=SUM(D2:D{total_row-1})'
                ws[f'E{total_row}'] = f'=SUM(E2:E{total_row-1})'
                ws[f'F{total_row}'] = f'=SUM(F2:F{total_row-1})'
                
                for col in ['D', 'E', 'F']:
                    ws[f'{col}{total_row}'].fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
                    ws[f'{col}{total_row}'].font = Font(bold=True)
            
            # Ширина колонок
            ws.column_dimensions['A'].width = 12
            ws.column_dimensions['B'].width = 25
            ws.column_dimensions['C'].width = 12
            ws.column_dimensions['D'].width = 12
            ws.column_dimensions['E'].width = 14
            ws.column_dimensions['F'].width = 12
            ws.column_dimensions['G'].width = 14
            
            # Центрирование числовых значений
            for row in ws.iter_rows(min_row=2, min_col=4, max_col=6, max_row=len(media)+1):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center')
            
            # Freeze pane
            ws.freeze_panes = 'A2'
            
            # Автофильтр
            if len(media) > 0:
                ws.auto_filter.ref = f'A1:G{len(media)+1}'
                
                # Диаграмма: топ медиа по лайкам
                top_n = min(15, len(media))
                chart = BarChart()
                chart.type = 'col'
                chart.title = f'Топ {top_n} медиа по лайкам'
                chart.y_axis.title = 'Лайки'
                chart.x_axis.title = 'Тип медиа'
                
                data = Reference(ws, min_col=4, min_row=1, max_row=top_n+1)
                cats = Reference(ws, min_col=3, min_row=2, max_row=top_n+1)
                chart.add_data(data, titles_from_data=True)
                chart.set_categories(cats)
                chart.height = 10
                chart.width = 18
                ws.add_chart(chart, 'A' + str(len(media) + 4))
            
            wb.save(filepath)
            logger.info("Медиа экспортировано в %s", filepath)
            return filepath
        except Exception as e:
            logger.error("Ошибка экспорта Excel: %s", e, exc_info=True)
            return None