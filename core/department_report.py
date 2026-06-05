"""Отчёт по кафедре за период (Word)."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

from core.database import Database


def export_department_report(
    department_id: int,
    date_from: str,
    date_to: str,
    dest_path: str | Path,
    db: Database | None = None,
) -> tuple[bool, str]:
    own = db is None
    if own:
        db = Database()
    try:
        dept = db.get_department_by_id(department_id)
        if not dept:
            return False, "Кафедра не найдена."
        posts = db.get_department_posts_in_period(department_id, date_from, date_to)
        employees = db.get_department_employees_activity(department_id, date_from, date_to)

        doc = Document()
        title = doc.add_heading(f"Отчёт: {dept.get('name', 'Кафедра')}", level=0)
        title.runs[0].font.size = Pt(16)
        doc.add_paragraph(f"Период: {date_from} — {date_to}")
        doc.add_paragraph(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph(f"Публикаций в архиве: {len(posts)}")

        doc.add_heading("Активность преподавателей", level=1)
        if employees:
            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "ФИО"
            hdr[1].text = "Хэштег"
            hdr[2].text = "Постов"
            for emp in employees:
                row = table.add_row().cells
                row[0].text = emp.get("full_name") or ""
                row[1].text = emp.get("hashtag") or ""
                row[2].text = str(emp.get("post_count", 0))
        else:
            doc.add_paragraph("Нет привязанных публикаций за период.")

        doc.add_heading("Публикации", level=1)
        for p in posts[:200]:
            doc.add_paragraph(
                f"#{p.get('original_post_id')} · {p.get('date', '')} · "
                f"лайки {p.get('likes', 0)} · { (p.get('text') or '')[:200] }",
                style="List Bullet",
            )
        if len(posts) > 200:
            doc.add_paragraph(f"… и ещё {len(posts) - 200} записей.")

        path = Path(dest_path)
        doc.save(str(path))
        return True, str(path)
    except Exception as e:
        return False, str(e)
    finally:
        if own and db:
            db.close()
